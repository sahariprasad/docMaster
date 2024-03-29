from docx import Document
import zipfile
import re
import sys

def stripper(longstring):
    longstring = longstring.replace("<bi:value><![CDATA[", "")
    longstring = longstring.replace(']]></bi:value>', '')
    longstring = longstring.lstrip()
    longstring = longstring.rstrip()
    return longstring

document = Document()
font = document.styles['Normal'].font
font.name = 'Calibri'

args = sys.argv
print(args[1])

lumxLocation = sys.argv[1]
lumxName = lumxLocation.split("\\")[-1].split('.')[0]
packageLocation = (lumxLocation.split(".")[0])
location2 = lumxLocation.split("\\")
location2.pop(-1)
location1 = ""
for parts in location2:
    location1 = location1 + parts + "\\"

outputXLSX = packageLocation + "\\" + lumxName + ".xlsx"
outputDOCX = location1 + "\\" + lumxName + ".docx"

document.add_heading(lumxName, 0)

with zipfile.ZipFile(lumxLocation, 'r') as zip_ref:
    zip_ref.extractall(packageLocation)

defaultApp = ""
headerFileLocation = packageLocation + "\\header.xml"
headerFile = open(headerFileLocation)

for line in headerFile:
    if re.search('<hi:property name="default_app">', line):
        defaultApp = line.split('default_app">')[1].split("<")[0]

fileLocation = packageLocation + "\\apps\\" + defaultApp + "\\content.biapp"

componentName = ""
dataSource = ""
dataSourceAliasArray = []
dataSourceTypeArray = []
dataSourceNameArray = []
componentArray = []
globalVariableArray = []
datasourceAlias = ""
datasourceName = ""
row = 1
col = 0

document.add_heading('Data Sources', level=1)
sourceFile = open(fileLocation, encoding="utf8")

for line in sourceFile:
    if re.search('bi:data_source_alias name=', line):
        ds_alias = line.split("name=\"")[1].split("\"")[0]
        dataSourceAliasArray.append(ds_alias)

    if re.search('bi:property name="DATA_SOURCE_TYPE"', line):
        ds_type = line.split("value=\"")[1].split("\"")[0]
        dataSourceTypeArray.append(ds_type)

    if re.search('bi:property name="DATA_SOURCE_NAME"', line):
        ds_name = line.split("value=\"")[1].split("\"")[0]
        dataSourceNameArray.append(ds_name)
        row = row + 1
table = document.add_table(rows=row, cols=3)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = "Data Source"
header_cells[1].text = "Data Source Type"
header_cells[2].text = "Data Source Name"

for x in range(row - 1):
    row_cells = table.rows[x + 1].cells
    row_cells[0].text = dataSourceAliasArray[x]
    row_cells[1].text = dataSourceTypeArray[x]
    row_cells[2].text = dataSourceNameArray[x]
document.add_heading('Component Data Source Mapping', level=1)
row = 1
col = 0
sourceFile = open(fileLocation, encoding="utf8")

for line2 in sourceFile:
    if re.search('<bi:component name=', line2):
        componentName = line2.split("name=\"")[1].split("\"")[0]
    if re.search('<bi:property name="DATA_SOURCE_ALIAS_REF"', line2):
        dataSource = line2.split("value=\"")[1].split("\"")[0]
        if dataSource != "":
            componentArray.append(componentName + " - " + dataSource)

            row = row + 1

table = document.add_table(rows=row, cols=2)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = "Component Name"
header_cells[1].text = "Mapped Data Source"

for x in range(row - 1):
    row_cells = table.rows[x + 1].cells
    row_cells[0].text = componentArray[x].split(' - ')[0]
    row_cells[1].text = componentArray[x].split(' - ')[1]
document.add_heading('Global Variables', level=1)
row = 0
col = 0
sourceFile = open(fileLocation, encoding="utf8")

for line in sourceFile:
    if re.search('bi:property name="GLOBALVARIABLE"', line):
        globalVarName = (next(sourceFile)).split("value=\"")[1].split("\"")[0]
        globalVariableArray.append(globalVarName)
        row = row + 1

table = document.add_table(rows=row, cols=2)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = "Global Variables"
header_cells[1].text = "Description"

for x in range(row - 1):
    row_cells = table.rows[x + 1].cells
    row_cells[0].text = globalVariableArray[x]
    row_cells[1].text = "Replace this text with description"
document.add_heading('Scripts', level=1)

# all components
sourceFile = open(fileLocation, encoding="utf8")
componentArray = []
canprintlines = False
biglineArray = []
bigline = ""
index = 0
for line2 in sourceFile:
    if re.search('<bi:component name=', line2):
        componentName = line2.split("name=\"")[1].split("\"")[0]
    elif re.search('bi:data_source_alias name=', line2):
        componentName = line2.split("name=\"")[1].split("\"")[0]
    elif re.search('<bi:property name="ON_STARTUP">', line2):
        componentName = "On Startup"
    elif re.search('bi:property name="ON_VAR_SUBMIT"', line2):
        componentName = "On Variable Initialization"
    elif re.search('<bi:property name="ON_BACKGROUND_PROCESSING">', line2):
        componentName = "On Background Processing"
    if re.search('<bi:property name="ON_', line2) or re.search('<bi:property name="on', line2):
        canprintlines = True
        document.add_heading(componentName, level=2)
    elif re.search('</bi:property>', line2):
        if canprintlines:
            bigline = stripper(bigline)
            document.add_paragraph(bigline)
            bigline = ""
        canprintlines = False
    if canprintlines:
        if line2.__contains__('<bi:property name="ON_') or line2.__contains__('<bi:property name="on'):
            0
        else:
            bigline = bigline + line2

document.add_heading('Global Script Objects', level=1)
sourceFile = open(fileLocation, encoding="utf8")
functionName = ""
bigline = ""
canprintlines = False
for line in sourceFile:
    if re.search('type="GLOBAL_SCRIPTS_COMPONENT"', line):
        functionName = (line.split("name=\"")[1].split("\"")[0])
        document.add_heading(functionName, level=2)
    if re.search('<bi:property name="FUNCTION"', line):
        # next(sourceFile)
        # print(line)
        document.add_heading((next(sourceFile).split('value="')[1].split('"')[0]), 3)
    if re.search('<bi:property name="FUNCTION_BODY">', line):
        canprintlines = True

    elif re.search('</bi:property>', line):
        if canprintlines:
            bigline = stripper(bigline)
            document.add_paragraph(bigline)
            bigline = ""
        canprintlines = False
    if canprintlines:
        if line.__contains__('<bi:property name="FUNCTION_BODY">'):
            0
        else:
            bigline = bigline + line

document.save(outputDOCX)
